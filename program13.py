# Program12
# Austin Riggs
# 11/22/23
import docx
import PyPDF2
import csv
import datetime
import json


print('\nRequirement 1\n')
print('This is Program13 - Austin Riggs')

print('\nRequirement 2\n')
print('This program demonstrates using Python to work with ' +
      'PDF and MS Word documents, and CSV and JSON data.')

print('\nRequirement 3\n')
print('Defining classes used for document processing...')
class PDFProcessing:
  def __init__(self, pdfFileName):
    self.pdfFileName = pdfFileName

  def getPageCount(self):
    pdfFile = open(self.pdfFileName, 'rb')
    pdfReader = PyPDF2.PdfReader(pdfFile)
    pageCount = len(pdfReader.pages)
    pdfFile.close()
    return pageCount

  def displayPage(self, pageNumber):
    pdfFile = open(self.pdfFileName, 'rb')
    pdfReader = PyPDF2.PdfReader(pdfFile)
    content = pdfReader.pages[pageNumber - 1].extract_text()
    pdfFile.close()
    return content


class WordProcessing:
  def __init__(self, wordFileName):
    self.wordFile = docx.Document(wordFileName)

  def getParagraphCount(self):
    return len(self.wordFile.paragraphs)

  def displayParagraph(self, paraNum):
    print('\n' + self.wordFile.paragraphs[paraNum - 1].text)


class CSVProcessing:
  def __init__(self, csvFileName):
    self.csvFileName = csvFileName

  def displayCSV(self):
    csvReadFile = open(self.csvFileName)
    csvReader = csv.DictReader(csvReadFile, ['time', 'fruit','qty'])
    csvContent = ''
    for row in csvReader:
      currentTime = row['time']
      currentFruit = row['fruit']
      currentQty = row['qty']
      csvContent += currentTime + ', ' + currentFruit + ', ' + currentQty + '\n'
    csvReadFile.close()
    return csvContent

  def updateCSV(self):
    csvWriteFile = open(self.csvFileName, 'a', newline='')
    csvWriter = csv.writer(csvWriteFile)

    newFruit = input("Please enter the name of a fruit to add to file: ")
    newQty = input("Please enter the quantity of fruit to add to file: ")
    dt = datetime.datetime.now()
    newTime= dt.strftime("%y-%m-%d %I:%M")

    csvWriter.writerow([newTime, newFruit, newQty])

    csvWriteFile.close()


class JSONProcessing:
  def __init__(self):
    self.dic = {}

  def inputData(self):
    print('Enter the names of 7 cities and an adjective describing the city.')
    for i in range(7):
      city = input('Enter the name of a city: ')
      adj = input('Enter an adjective that describes the city: ')
      self.dic[city] = adj

  def convertToJSON(self):
    jsonData = json.dumps(self.dic)
    return jsonData



print('\nRequirement 4\n')
pdfObj = PDFProcessing('meetingminutes.pdf')
print("PDF document page count: ", pdfObj.getPageCount())

print('\nRequirement 5\n')
pageNum = int(input('Please enter the page number of the '
                    'pdf file you want to see: '))
pdf_string = pdfObj.displayPage(pageNum)
print('Content of PDF page: ')
print('\n' + pdf_string)

print('\nRequirement 6\n')
docObj = WordProcessing('demo.docx')
print("Word document paragraph count: ", docObj.getParagraphCount())

print('\nRequirement 7\n')
paragraphNum = int(input('Enter the paragraph number of the '
                         'document you want to see: '))
print('Paragraph content:')
docObj.displayParagraph(paragraphNum)

print('\nRequirement 8\n')
csvObj = CSVProcessing('example.csv')
csvString = csvObj.displayCSV()
print("Content of CSV File:")
print(csvString)

print('\nRequirement 9\n')
csvObj.updateCSV()
print("\nContent of updates CVS file:")
csvUpdated = csvObj.displayCSV()
print(csvUpdated)

print('\nRequirement 10\n')
jsonObj = JSONProcessing()
jsonObj.inputData()

print('\nRequirement 11\n')
print('JSON-formatted data:')
jsonFormat = jsonObj.convertToJSON()
print(jsonFormat)

print('\nRequirement 12\n')
print('''I learned a lot from this assignment. Similarly to the last 
assignment, I have not programmatically maniplated documents in the past
and learned how to do so with many document formats here. I will definitely 
be using some of the modules in the future. Looking forward to the last 
assignment.''')





